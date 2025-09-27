from docx import Document
from datetime import datetime
import os

def gerar_ata(titulares, suplentes, cidade_processo):
    modelo = "MODELO ATA EXPORTAR.docx"
    if not os.path.exists(modelo):
        raise FileNotFoundError("❌ Modelo de ata não encontrado. Verifique o caminho.")

    doc = Document(modelo)

    # Data e hora
    hoje = datetime.now()
    dia = str(hoje.day)
    # mes = hoje.strftime("%B")
    meses_pt = {
    "January": "janeiro", "February": "fevereiro", "March": "março",
    "April": "abril", "May": "maio", "June": "junho",
    "July": "julho", "August": "agosto", "September": "setembro",
    "October": "outubro", "November": "novembro", "December": "dezembro"
    }
    mes = meses_pt[hoje.strftime("%B")]
    ano = str(hoje.year)
    hora = hoje.strftime("%H:%M")
    
    


    # Substituir campos de data/hora
    for p in doc.paragraphs:
        p.text = p.text.replace("[informara o dia]", dia)
        p.text = p.text.replace("[informara o mês]", mes)
        p.text = p.text.replace("[informara o ano]", ano)
        p.text = p.text.replace("[informara a hora]", hora)

    # Substituir o marcador pela lista formatada
    for i, p in enumerate(doc.paragraphs):
        if "[a lista de jurados e suplentes ficará aqui]" in p.text:
            recuo = p.paragraph_format.left_indent
            parent = p._element.getparent()
            idx = parent.index(p._element)
            parent.remove(p._element)

            pos = idx

            # Titulares: 1 a 25
            for idx_j, j in enumerate(titulares, 1):
                prof = getattr(j, "profissao", None)
                nome_fmt = j.nome.upper() + (f" ({prof.upper()})" if prof else "")
                par = doc.add_paragraph(
                    f"{idx_j}. {nome_fmt}, {j.endereco}, {j.numero}, {j.bairro}, {j.cidade};"
                )
                par.paragraph_format.left_indent = recuo
                parent.insert(pos, par._element)
                pos += 1

            # Adicionar separador "SUPLENTES:"
            par_suplentes = doc.add_paragraph("SUPLENTES:")
            par_suplentes.paragraph_format.left_indent = recuo
            parent.insert(pos, par_suplentes._element)
            pos += 1

            # Suplentes: 1 a 5
            for idx_j, j in enumerate(suplentes, 1):
                prof = getattr(j, "profissao", None)
                nome_fmt = j.nome.upper() + (f" ({prof.upper()})" if prof else "")
                par = doc.add_paragraph(
                    f"{idx_j}. {nome_fmt}, {j.endereco}, {j.numero}, {j.bairro}, {j.cidade};"
                )
                par.paragraph_format.left_indent = recuo
                parent.insert(pos, par._element)
                pos += 1

            break

    nome_arquivo = f"ata_sorteio_{cidade_processo.lower()}_{hoje.strftime('%Y%m%d_%H%M')}.docx"
    doc.save(nome_arquivo)
    print(f"✅ Ata final gerada com separação de suplentes: {nome_arquivo}")
    return nome_arquivo








# def gerar_ata(titulares, suplentes, cidade_processo):
#     caminho_modelo = os.path.join(os.path.dirname(__file__), "MODELO ATA EXPORTAR.docx")
#     doc = Document(caminho_modelo)

#     # Traduzir mês para português
#     hoje = datetime.now()
#     meses = {
#         "January": "janeiro", "February": "fevereiro", "March": "março",
#         "April": "abril", "May": "maio", "June": "junho",
#         "July": "julho", "August": "agosto", "September": "setembro",
#         "October": "outubro", "November": "novembro", "December": "dezembro"
#     }
#     mes_en = hoje.strftime("%B")
#     mes = meses.get(mes_en, mes_en)

#     dia = hoje.day
#     ano = hoje.year
#     hora = hoje.strftime("%H:%M")

#     # Substituir os campos da ata
#     for p in doc.paragraphs:
#         if "[informara o dia]" in p.text:
#             p.text = p.text.replace("[informara o dia]", str(dia))
#             p.text = p.text.replace("[informara o mês]", mes)
#             p.text = p.text.replace("[informara o ano]", str(ano))
#             p.text = p.text.replace("[informara a hora]", hora)
#             break

#     # Adicionar jurados titulares
#     doc.add_paragraph("JURADOS SORTEADOS:\n")

#     for i, j in enumerate(titulares, 1):
#         doc.add_paragraph(f"{i}. {j.nome.upper()}, {j.endereco}, {j.numero}, {j.bairro};")

#     # Adicionar suplentes
#     doc.add_paragraph("\nSUPLENTES:\n")

#     for i, j in enumerate(suplentes, 1):
#         doc.add_paragraph(f"{i}. {j.nome.upper()}, {j.endereco}, {j.numero}, {j.bairro};")

#     # Salvar arquivo final
#     nome_arquivo = f"ata_sorteio_{cidade_processo.lower()}_{hoje.strftime('%Y%m%d_%H%M')}.docx"
#     doc.save(nome_arquivo)
#     print(f"✅ Ata gerada: {nome_arquivo}")
#     return nome_arquivo
